from dotenv import load_dotenv
load_dotenv()

from collections import defaultdict
from docx import Document

from models.interaction_model import get_all_user_interactions

# ====================================
# AMBIL SEMUA DATA INTERAKSI
# ====================================
data = get_all_user_interactions()

# ====================================
# KELOMPOKKAN BERDASARKAN SESSION_ID
# ====================================
sessions = defaultdict(list)

for item in data:

    session_id = item.get("session_id")

    if not session_id:
        continue

    sessions[session_id].append(item)

# ====================================
# BUAT DOKUMEN WORD
# ====================================
doc = Document()

doc.add_heading("Hasil Evaluasi Precision@N", level=1)

# ====================================
# TABEL HASIL EVALUASI
# ====================================
table = doc.add_table(rows=1, cols=6)
table.style = "Table Grid"

hdr_cells = table.rows[0].cells

hdr_cells[0].text = "Session ID"
hdr_cells[1].text = "User ID"
hdr_cells[2].text = "Query"
hdr_cells[3].text = "Jumlah Like"
hdr_cells[4].text = "Total Rekomendasi"
hdr_cells[5].text = "Precision@N"

# ====================================
# HITUNG PRECISION
# ====================================
total_precision = 0
jumlah_session = 0

print("\n====================================")
print("HASIL EVALUASI PRECISION@N")
print("====================================")

for session_id, interactions in sessions.items():

    # ====================================
    # JUMLAH ITEM RELEVAN
    # ====================================
    relevant_count = len(interactions)

    # ====================================
    # TOTAL REKOMENDASI
    # ====================================
    total_result = interactions[0].get("total_result", 0)

    if total_result == 0:
        continue

    # ====================================
    # HITUNG PRECISION
    # ====================================
    precision = relevant_count / total_result

    total_precision += precision
    jumlah_session += 1

    # ====================================
    # AMBIL DATA TAMBAHAN
    # ====================================
    user_id = interactions[0].get("user_id")
    query = interactions[0].get("query")

    # ====================================
    # TAMPILKAN DI TERMINAL
    # ====================================
    print(f"\nSession ID        : {session_id}")
    print(f"User ID           : {user_id}")
    print(f"Query             : {query}")
    print(f"Jumlah Like       : {relevant_count}")
    print(f"Total Rekomendasi : {total_result}")
    print(f"Precision@N       : {precision:.2f}")

    # ====================================
    # TAMBAH KE WORD
    # ====================================
    row_cells = table.add_row().cells

    row_cells[0].text = str(session_id)
    row_cells[1].text = str(user_id)
    row_cells[2].text = str(query)
    row_cells[3].text = str(relevant_count)
    row_cells[4].text = str(total_result)
    row_cells[5].text = f"{precision:.2f}"

# ====================================
# RATA-RATA PRECISION
# ====================================
if jumlah_session > 0:

    average_precision = total_precision / jumlah_session

    print("\n====================================")
    print(f"RATA-RATA PRECISION@N : {average_precision:.2f}")
    print("====================================")

    # ====================================
    # TULIS RATA-RATA KE WORD
    # ====================================
    p = doc.add_paragraph()
    p.add_run("RATA-RATA PRECISION@N : ").bold = True
    p.add_run(f"{average_precision:.2f}")

else:

    print("\nTidak ada data interaksi.")

    doc.add_paragraph("Tidak ada data interaksi.")

# ====================================
# SIMPAN FILE WORD
# ====================================
doc.save("evaluasi.docx")

print("\nSELESAI!")
print("File tersimpan: evaluasi.docx")